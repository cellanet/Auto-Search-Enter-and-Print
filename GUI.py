#Gia Minh Hoang Jan 27 2022
# Auto search for customer name, upload selected info and print out a hard copy

import os
from tkinter.ttk import Combobox
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from selenium.webdriver import ActionChains
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
import time

import tkinter as tk
from tkinter import *

import Open_ERP
from Open_ERP import driver
import pyperclip as pyclip


def sort_package():
    try:

        click_link = driver.find_element(By.LINK_TEXT, "Customer Service")
        click_link.click()
        time.sleep(1.5)
        search = driver.find_element(By.CLASS_NAME, "oe_searchview_input")
        search.click()

        search.clear()
        search.send_keys(ticket_box.get())
        search.send_keys(Keys.RETURN)

        time.sleep(1.5)
        found = driver.find_element(By.XPATH, "//td[@title='Number of customer service ticket']")
        found.click()

        time.sleep(1.5)
        check0 = driver.find_element(By.CSS_SELECTOR, "span[class='oe_form_uri']").text

        # if check0 = refund, check refund item
        if check0 == 'Refund Requests':
            check1 = driver.find_element(By.XPATH,
                                         "//span[@class='oe_form_uri' and not(contains(text(),'Refund Requests'))]").text
        # if check0 = service, check service item
        elif check0 == 'Services/Repairs':
            check1 = driver.find_element(By.XPATH,
                                         "//span[@class='oe_form_uri' and not(contains(text(),'Services/Repairs'))]").text
        elif check0 == 'Warranty/Replacements':
            check1 = driver.find_element(By.XPATH,
                                         "//span[@class='oe_form_uri' and not(contains(text(),'Warranty/Replacements'))]").text

        ticket_note = ''
        #  Battery, Light, System
        if check1 != 'Loupes' and check1 != 'Loaner- Complimentary':
            ticket_note = 'Received, gave to ****.'
            ticket_box.delete(0, END)

        # Loupes
        elif check1 == 'Loupes' or check1 == 'Loaner- Complimentary':
            ticket_note = 'Received, gave to ****'
            loupes()

        if ticket_note != '':
            lognote = driver.find_element(By.XPATH, "//a[@class= 'oe_compose_log']")
            action = ActionChains(driver)
            action.click(lognote)
            action.pause(0.5)
            action.send_keys(ticket_note)
            action.perform()

            click_to_post = driver.find_element(By.XPATH, "//button[@class = 'oe_post']")
            click_to_post.click()
    except:
        print("error1")


root = tk.Tk()
root.title('Service')
root.iconbitmap('C:/Users/****/OneDrive/Desktop/SERVICE/setting.ico')
canvas = tk.Canvas(root, width=300, height=80)
root.resizable(False, False)
canvas.grid(columnspan=3, rowspan=5)
root.grid_columnconfigure(0, weight=0)
root.grid_rowconfigure(0, weight=0)

# ========================== Find ticket box =============================

ticket_frame = LabelFrame(root, text='Find Ticket')  # FRAME
ticket_frame.grid(column=0, row=0, padx=5)

input_label = tk.Label(ticket_frame, text="Enter ticket number: ")  # LABEL
input_label.grid(column=0, row=0, pady=10, padx=5)

ticket_box = tk.Entry(ticket_frame, width=20)  # INPUT BOX
ticket_box.grid(column=1, row=0, padx=(0, 50))

enter_button = Button(ticket_frame, text='1.Find', activebackground='#b7e4c7', command=sort_package)  # BUTTON
enter_button.grid(column=0, row=1, padx=(0, 0), pady=(5, 10))

# address_check = IntVar()
# address = Checkbutton(ticket_frame, text='Address Confirmed', variable=address_check)
# address.grid(column=1, row=1, padx=(0, 87))

address_check = StringVar()
address = Combobox(ticket_frame, width=15, textvariable=address_check)
address['values'] = ('', 'No Service Form', 'Address confirmed', 'Address Updated')
address.grid(column=1, row=1, padx=(0, 87))
# address.current()


# ================================================================
# =========================== Misc Box ===========================
misc_frame = LabelFrame(root, text='Miscellaneous')  # light Frame
misc_frame.grid(column=1, row=0, padx=5)

loupes_check = IntVar()
loupes = Checkbutton(misc_frame, text='Loupes', variable=loupes_check)  # Loupes
loupes.grid(column=0, row=1)

charger_check = IntVar()
charger = Checkbutton(misc_frame, text='Charger', variable=charger_check)  # charger
charger.grid(column=1, row=1)

comp_check = IntVar()
comp = Checkbutton(misc_frame, text='Comp filter', variable=comp_check)  # comp filter
comp.grid(column=2, row=1, padx=(0, 47))

else_misc_label = Label(misc_frame, text='Else:')
else_misc_label.grid(column=0, row=4)
else_misc_box = Text(misc_frame, width=15, height=3)
else_misc_box.grid(column=1, row=4, padx=10, pady=5)

# ================================================================
# =========================== Battery Box ========================

battery_frame = LabelFrame(root, text='Battery')  # Frame
battery_frame.grid(column=0, row=2, padx=5, pady=5)

serial_label = Label(battery_frame, text='Type/Serial:')  # Serial Label
serial_label.grid(column=0, row=0)

combobox_num = StringVar()
battery_type1 = Combobox(battery_frame, width=10, textvariable=combobox_num)
battery_type1['values'] = ('', '3rd Gen O', '2nd Gen O', '1st Gen O', 'AHF', 'ATHF', 'HF', 'THF', 'V3 Battery ')
battery_type1.grid(column=1, row=0, padx=(0, 150))
battery_type1.current()

serial_box = Entry(battery_frame, width=7)  # battery's serial box
serial_box.grid(column=1, row=0)

working_check = IntVar()
working = Checkbutton(battery_frame, text='Not working', variable=working_check)  # check not working
working.grid(column=0, row=1, padx=(16, 0))

work_check = IntVar()
work = Checkbutton(battery_frame, text='Working', variable=work_check)  # check working
work.grid(column=1, row=1, padx=(16, 0))

solder_check = IntVar()
solder = Checkbutton(battery_frame, text='Re-solder', variable=solder_check)  # check solder
solder.grid(column=0, row=2)

hour_check = IntVar()
hour = Checkbutton(battery_frame, text='Case cracked', variable=hour_check)  # check lasting hours
hour.grid(column=1, row=2, padx=(42, 0))

else_label = Label(battery_frame, text='Else:')
else_label.grid(column=0, row=4)
else_box = Text(battery_frame, width=15, height=3)
else_box.grid(column=1, row=4, padx=30, pady=(5, 10))

# ================================================================

# =========================== Battery2 Box =======================
battery_frame2 = LabelFrame(root, text='Battery 2')  # Frame
battery_frame2.grid(column=1, row=2, padx=5, pady=5)

serial_label2 = Label(battery_frame2, text='Type/Serial:')  # Serial Label
serial_label2.grid(column=0, row=0)

combobox_num2 = StringVar()
battery_type2 = Combobox(battery_frame2, width=10, textvariable=combobox_num2)
battery_type2['values'] = ('', '3rd Gen O', '2nd Gen O', '1st Gen O', 'AHF', 'ATHF', 'HF', 'THF', 'V3 Battery ')
battery_type2.grid(column=1, row=0, padx=(0, 150))
battery_type2.current()

serial_box2 = Entry(battery_frame2, width=7)  # battery's serial box
serial_box2.grid(column=1, row=0)

working_check2 = IntVar()
working2 = Checkbutton(battery_frame2, text='Not working', variable=working_check2)  # check not working
working2.grid(column=0, row=1, padx=(16, 0))

work_check2 = IntVar()
work2 = Checkbutton(battery_frame2, text='Working', variable=work_check2)  # check working
work2.grid(column=1, row=1, padx=(16, 0))

solder_check2 = IntVar()
solder2 = Checkbutton(battery_frame2, text='Re-solder', variable=solder_check2)  # check solder
solder2.grid(column=0, row=2)

hour_check2 = IntVar()
hour2 = Checkbutton(battery_frame2, text='Case cracked', variable=hour_check2)  # check lasting hours
hour2.grid(column=1, row=2, padx=(42, 0))

else_label2 = Label(battery_frame2, text='Else:')
else_label2.grid(column=0, row=4)
else_box2 = Text(battery_frame2, width=15, height=3)
else_box2.grid(column=1, row=4, padx=30, pady=(5, 10))

# ================================================================

# =========================== Light Box ==========================
light_frame = LabelFrame(root, text='Light')  # light Frame
light_frame.grid(column=0, row=3, padx=5, pady=5)

light_label = Label(light_frame, text='Type/Serial:')  # Serial Label
light_label.grid(column=0, row=0)

combobox_light1 = StringVar()
light_type1 = Combobox(light_frame, width=10, textvariable=combobox_light1)
light_type1['values'] = ('', '3rd Gen Light', 'Elite', 'Cable', 'V3 Light ')
light_type1.grid(column=1, row=0, padx=(0, 150))
light_type1.current()

light_box = Entry(light_frame, width=7)  # light's serial box
light_box.grid(column=1, row=0)

light_work_check = IntVar()
light_work = Checkbutton(light_frame, text='Working', variable=light_work_check)  # check working
light_work.grid(column=1, row=1, padx=(16, 0))

plug_check = IntVar()
plug = Checkbutton(light_frame, text='Plug worn/damaged', variable=plug_check)  # check plug
plug.grid(column=0, row=2, padx=(8, 0))

cord_check = IntVar()
cord = Checkbutton(light_frame, text='Hard cord', variable=cord_check)  # check lasting hours
cord.grid(column=1, row=2, padx=(23, 0))

led_check = IntVar()
led = Checkbutton(light_frame, text='LED fell off', variable=led_check)  # check working
led.grid(column=0, row=1, padx=(0, 43))

else_light_label = Label(light_frame, text='Else:')
else_light_label.grid(column=0, row=4)
else_light_box = Text(light_frame, width=15, height=3)
else_light_box.grid(column=1, row=4, padx=0, pady=(0, 10))
# ================================================================
# =========================== Light 2 Box ========================
light_frame2 = LabelFrame(root, text='Light 2')  # light Frame
light_frame2.grid(column=1, row=3, padx=5, pady=5)

light_label2 = Label(light_frame2, text='Type/Serial:')  # Serial Label
light_label2.grid(column=0, row=0)

combobox_light2 = StringVar()
light_type2 = Combobox(light_frame2, width=10, textvariable=combobox_light2)
light_type2['values'] = ('', '3rd Gen Light', 'Elite', 'Cable', 'V3 Light ')
light_type2.grid(column=1, row=0, padx=(0, 150))
light_type2.current()

light_box2 = Entry(light_frame2, width=7)  # light's serial box
light_box2.grid(column=1, row=0)

light_work_check2 = IntVar()
light_work2 = Checkbutton(light_frame2, text='Working', variable=light_work_check2)  # check working
light_work2.grid(column=1, row=1, padx=(16, 0))

plug_check2 = IntVar()
plug2 = Checkbutton(light_frame2, text='Plug worn/damaged', variable=plug_check2)  # check plug
plug2.grid(column=0, row=2, padx=(8, 0))

cord_check2 = IntVar()
cord2 = Checkbutton(light_frame2, text='Hard cord', variable=cord_check2)  # check lasting hours
cord2.grid(column=1, row=2, padx=(23, 0))

led_check2 = IntVar()
led2 = Checkbutton(light_frame2, text='LED fell off', variable=led_check2)  # check working
led2.grid(column=0, row=1, padx=(0, 43))

else_light_label2 = Label(light_frame2, text='Else:')
else_light_label2.grid(column=0, row=4)
else_light_box2 = Text(light_frame2, width=15, height=3)
else_light_box2.grid(column=1, row=4, padx=30, pady=(0, 10))


# ================================= Functions =====================================

def loupes():
    document = Document()
    section = document.sections[0]
    section.page_height = Inches(6)
    section.page_width = Inches(4)
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(9)

    style = document.styles.add_style('Name1', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.bold = True
    font.underline = True

    style = document.styles.add_style('Name2', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(60)

    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.1)
        section.bottom_margin = Inches(0.1)
        section.left_margin = Inches(0.2)
        section.right_margin = Inches(0.2)

    check0 = driver.find_element(By.CSS_SELECTOR, "span[class='oe_form_uri']").text
    check1 = ''
    if check0 == 'Refund Requests':
        check1 = driver.find_element(By.XPATH,
                                     "//span[@class='oe_form_uri' and not(contains(text(),'Refund Requests'))]").text
    elif check0 == 'Services/Repairs':
        check1 = driver.find_element(By.XPATH,
                                     "//span[@class='oe_form_uri' and not(contains(text(),'Services/Repairs'))]").text
    elif check0 == 'Warranty/Replacements':
        check1 = driver.find_element(By.XPATH,
                                     "//span[@class='oe_form_uri' and not(contains(text(),'Warranty/Replacements'))]").text

    # ticketnumber_input =
    document.add_paragraph('\n  ' + ticket_box.get(), style='Name2')
    document.add_paragraph(check0 + ': ' + check1, style='Name1')

    document.save(r'C:\Users\****\OneDrive\Desktop\ERVICE\test.doc')
    os.startfile("C:/Users/****/OneDrive/Desktop/SERVICE/test.doc", "print")

def create_doc():
    document = Document()
    section = document.sections[0]
    section.page_height = Inches(6)
    section.page_width = Inches(4)
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(9)

    style = document.styles.add_style('Name1', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.bold = True
    font.underline = True

    style = document.styles.add_style('Name2', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(60)

    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.1)
        section.bottom_margin = Inches(0.1)
        section.left_margin = Inches(0.2)
        section.right_margin = Inches(0.2)

    # def delete_paragraph(paragraph):
    #     p = paragraph._element
    #     p.getparent().remove(p)
    #     p._p = p._element = None

    # String
    ticketnumber_text = 'Ticket: ' + ticket_box.get()
    battery_text = '*Cell date\n*Working\n*Lasting hrs\n*Loose cnx\n*Charger plug\n*ON OFF\n*(+)(-)\n*Case\n*Clip\n   FIXED:\n\n'
    light_text = '*Working\n*Cable\n  -Hard\n  -Working\n*LED\n*Pins\n*Housing\n   FIXED:\n\n'

    # =======# =======# =======# =======# =======# =======# =======# =======# =======# =======# =======# =======

    # create table 1 by 3
    table = document.add_table(1, 3, style='Table Grid')
    table.autofit = True

    for cell in table.columns[0].cells:
        cell.width = Inches(2)

    # ===================== Battery1 =========================
    # if have battery
    battery1 = battery_type1.get() + serial_box.get()
    if (working_check.get() == 1):
        battery1 = battery1 + ': Not working'
    if (work_check.get() == 1):
        battery1 = battery1 + ': Working'
    if (solder_check.get() == 1):
        battery1 = battery1 + '\n  *Need to be re-solder.'
    if (hour_check.get() == 1):
        battery1 = battery1 + '\n  *Case cracked.'
    if (len(else_box.get("1.0", "end-1c"))):
        battery1 = battery1 + '\n  *' + else_box.get("1.0", "end-1c")

    # ===================== Battery2 =========================
    # battery2 = ''
    # if len(serial_box2.get()):
    #     battery2 = '\n' + battery_type2.get() + serial_box2.get()
    battery2 = '\n' + battery_type2.get() + serial_box2.get()
    if (working_check2.get() == 1):
        battery2 = battery2 + ': Not working'
    if (work_check2.get() == 1):
        battery2 = battery2 + ': Working'
    if (solder_check2.get() == 1):
        battery2 = battery2 + '\n  *Need to be re-solder.'
    if (hour_check2.get() == 1):
        battery2 = battery2 + '\n  *Case cracked.'
    if (len(else_box2.get("1.0", "end-1c"))):
        battery2 = battery2 + '\n  *' + else_box2.get("1.0", "end-1c")

    # ===================== Light 1 =========================

    # if have light
    # light1 = ''
    # if len(light_box.get()):
    #     light1 = '\n' + light_type1.get() + light_box.get()

    light1 = '\n' + light_type1.get() + light_box.get()
    if light_work_check.get() == 1:
        light1 = light1 + ' Working'
    if led_check.get() == 1:
        light1 = light1 + '\n  *LED fell off'
    if plug_check.get() == 1:
        light1 = light1 + '\n  *Plug worn/damaged'
    if cord_check.get() == 1:
        light1 = light1 + '\n  *Hard cord'
    if len(else_light_box.get("1.0", "end-1c")):
        light1 = light1 + '\n  *' + else_light_box.get("1.0", "end-1c")

    # ===================== Light 2 =========================
    # if have light
    # light2 = ''
    # if len(light_box2.get()):
    #     light2 = '\n' + light_type2.get() + light_box2.get()
    light2 = '\n' + light_type2.get() + light_box2.get()

    if light_work_check2.get() == 1:
        light2 = light2 + ' Working'
    if led_check2.get() == 1:
        light2 = light2 + '\n  *LED fell off'
    if plug_check2.get() == 1:
        light2 = light2 + '\n  *Plug worn/damaged'
    if cord_check2.get() == 1:
        light2 = light2 + '\n  *Hard cord'
    if len(else_light_box2.get("1.0", "end-1c")):
        light2 = light2 + '\n  *' + else_light_box2.get("1.0", "end-1c")

    # ===================== Misc =========================
    # everything else
    misc1 = '\n'
    if comp_check.get() == 1:
        misc1 = misc1 + '\nComp filter'
    if charger_check.get() == 1:
        misc1 = misc1 + '\nCharger'
    if (loupes_check.get() == 1):
        misc1 = misc1 + '\nLoupes'
    if (len(else_misc_box.get("1.0", "end-1c"))):
        misc1 = misc1 + '\n' + else_misc_box.get("1.0", "end-1c")

    # ===================== DONE GUI =========================

    check_text = '\nCheck:\n\n' + battery1 + battery2 + light1 + light2 + misc1  # copy everything
    pyclip.copy(check_text)

    ticketnumber = table.cell(0, 0).add_paragraph(ticketnumber_text, style='Name1')
    ticketnumber = table.cell(0, 0).add_paragraph(check_text)

    # battery
    battery = table.cell(0, 1).add_paragraph('Battery: ', style='Name1')
    battery = table.cell(0, 1).add_paragraph(battery_text)

    # light
    light = table.cell(0, 2).add_paragraph('Light: ', style='Name1')
    light = table.cell(0, 2).add_paragraph(light_text)

    check0 = driver.find_element(By.CSS_SELECTOR, "span[class='oe_form_uri']").text
    check1 = ''

    if check0 == 'Refund Requests':
        check1 = driver.find_element(By.XPATH,
                                     "//span[@class='oe_form_uri' and not(contains(text(),'Refund Requests'))]").text
    elif check0 == 'Services/Repairs':
        check1 = driver.find_element(By.XPATH,
                                     "//span[@class='oe_form_uri' and not(contains(text(),'Services/Repairs'))]").text

    # ticketnumber_input =
    document.add_paragraph('\n  ' + ticket_box.get(), style='Name2')
    document.add_paragraph(check0 + ': ' + check1, style='Name1')

    # change path here
    document.save(r'C:\Users\****\OneDrive\Desktop\SERVICE\test.doc')
    os.startfile("C:/Users/****/OneDrive/Desktop/SERVICE/test.doc", "print")


def paste_check():
    if address.get() != '':
        lognote1 = driver.find_element(By.XPATH, "//a[@class= 'oe_compose_log']")
        # paste value
        action = ActionChains(driver)
        action.click(lognote1)
        action.pause(0.5)
        action.send_keys(address.get())
        action.perform()
        click_to_post = driver.find_element(By.XPATH, "//button[@class = 'oe_post']")
        click_to_post.click()

        # hand off ticket
    time.sleep(1.5)
    handoff = driver.find_element(By.XPATH, "//button[@class= 'oe_button oe_form_button oe_highlight']")
    handoff.click()
    time.sleep(1)
    click_link = driver.find_element(By.LINK_TEXT, "Cancel")
    click_link.click()

    # click log a note
    lognote = driver.find_element(By.XPATH, "//a[@class= 'oe_compose_log']")

    # paste value
    action = ActionChains(driver)
    action.click(lognote)
    action.pause(0.5)
    action.key_down(Keys.CONTROL)
    action.send_keys('v')
    action.key_up(Keys.CONTROL)
    action.perform()


def submit_button():
    create_doc()
    paste_check()


def clear_all():
    serial_box.delete(0, END)
    serial_box2.delete(0, END)
    working_check.set(0)
    working_check2.set(0)
    work_check.set(0)
    work_check2.set(0)
    solder_check.set(0)
    solder_check2.set(0)
    hour_check.set(0)
    hour_check2.set(0)
    else_box.delete("1.0", END)
    else_box2.delete("1.0", END)
    light_box.delete(0, END)
    light_box2.delete(0, END)
    light_work_check.set(0)
    light_work_check2.set(0)
    plug_check.set(0)
    plug_check2.set(0)
    cord_check.set(0)
    cord_check2.set(0)
    led_check.set(0)
    led_check2.set(0)
    else_light_box.delete("1.0", END)
    else_light_box2.delete("1.0", END)
    ticket_box.delete(0, END)
    loupes_check.set(0)
    charger_check.set(0)
    comp_check.set(0)
    else_misc_box.delete("1.0", END)
    address_check.set('')
    battery_type1.set('')
    battery_type2.set('')
    light_type1.set('')
    light_type2.set('')
    ticket_box.focus()


# ============================ BUTTONS =====================

submit_button = Button(root, text='2.Submit',
                       activebackground='#b7e4c7', command=submit_button)
submit_button.grid(column=0, row=5, pady=10)

delete_button = Button(root, text='  3.Clear  ', activebackground='#b7e4c7', command=clear_all)
delete_button.grid(column=1, row=5, pady=10)

# ================================================================

# def hit_enter(e):
#     sort_package()
#
#
# root.bind('<Return>', hit_enter)

canvas = tk.Canvas(root, width=10, height=10)
canvas.grid(columnspan=3)
root.mainloop()
driver.quit()
